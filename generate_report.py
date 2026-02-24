from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Title
title = doc.add_heading('Rapport Technique - InvestBuddy', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtitle
subtitle = doc.add_paragraph('Projet DATA - Architecture de l\'application')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# Section 1: Architecture du systeme
doc.add_heading('1. Architecture du systeme', level=1)

doc.add_paragraph(
    "InvestBuddy est une application web reposant sur une architecture microservices conteneurisee avec Docker Compose. "
    "Cette architecture permet de separer les differentes responsabilites de l'application en services independants qui communiquent "
    "entre eux via un reseau Docker dedie. L'application est composee de cinq services principaux : une interface web, une API agent, "
    "une API de prediction, un orchestrateur n8n et une base de donnees MySQL."
)

doc.add_paragraph(
    "Le reseau Docker utilise est nomme investbuddy_net et fonctionne avec le driver bridge. Ce reseau permet a tous les services "
    "de communiquer entre eux en utilisant leurs noms de service comme noms d'hote. Par exemple, l'interface web peut appeler l'API agent "
    "en utilisant l'adresse http://agent-api:4000. Cette approche isole l'application du reste du systeme hote tout en permettant "
    "une communication fluide entre les composants."
)

# Section 2: Description des services
doc.add_heading('2. Description des services', level=1)

doc.add_heading('2.1 Service Webapp', level=2)
doc.add_paragraph(
    "Le service webapp constitue l'interface utilisateur de l'application. Il est developpe avec React 18, TypeScript et Vite, "
    "ce qui permet une experience utilisateur reactive et moderne. Ce service est expose sur le port 5173 et offre trois fonctionnalites "
    "principales : un tableau de bord, un assistant IA interactif et un glossaire de termes financiers. L'interface communique avec "
    "l'API agent pour les questions posees a l'assistant et avec l'API de prediction pour obtenir des previsions de volatilite."
)

doc.add_heading('2.2 Service Agent API', level=2)
doc.add_paragraph(
    "L'API agent est developpee avec FastAPI et Python 3.10. Elle est exposee sur le port 4000 et assure deux roles essentiels. "
    "Le premier role est l'agent pedagogique base sur une architecture RAG (Retrieval Augmented Generation) qui permet de repondre "
    "aux questions des utilisateurs en utilisant une base de connaissances financieres vectorisee. Le deuxieme role est l'analyse "
    "de risques qui combine les predictions du modele ML avec des indicateurs de marche temps reel comme le Fear and Greed Index "
    "et le Funding Rate de Binance. Cette API utilise ChromaDB comme base de donnees vectorielle et les embeddings d'OpenAI "
    "pour la recherche semantique."
)

doc.add_heading('2.3 Service Prediction API', level=2)
doc.add_paragraph(
    "L'API de prediction est egalement developpee avec FastAPI et est exposee sur le port 8080. Elle constitue le coeur du systeme ML "
    "en assurant le serving des modeles de prediction de volatilite. Cette API expose trois endpoints principaux. L'endpoint /predict "
    "permet d'obtenir une prediction de volatilite et de direction pour une cryptomonnaie donnee. L'endpoint /notify declenche "
    "l'envoi d'une notification via n8n. L'endpoint /feedback permet de collecter les retours utilisateurs pour ameliorer le modele. "
    "Les modeles utilises sont des RandomForestRegressor entraînes pour predire la volatilite et la direction du marche. "
    "Ces modeles sont stockes sous forme de fichiers pickle dans le dossier artifacts."
)

doc.add_heading('2.4 Service n8n', level=2)
doc.add_paragraph(
    "Le service n8n est un orchestrateur de workflows base sur l'image officielle n8nio/n8n. Il est expose sur le port 5678 "
    "et permet d'automatiser les notifications et la boucle Human-in-the-Loop. Le workflow implemente, nomme investbuddy_alert.json, "
    "fonctionne de la maniere suivante. Lorsqu'une prediction avec une volatilite elevee est detectee, n8n genere automatiquement "
    "un message personnalise grace a GPT-3.5 et envoie un email a l'utilisateur. Cet email contient deux liens permettant a l'utilisateur "
    "de confirmer ou d'infirmer la prediction, ce qui constitue le mecanisme de feedback Human-in-the-Loop."
)

doc.add_heading('2.5 Service MySQL', level=2)
doc.add_paragraph(
    "Le service MySQL utilise l'image officielle mysql:8.0 et est expose sur le port 3306. Il assure la persistance des donnees "
    "de l'application, notamment les informations utilisateurs et les historiques de predictions. La base de donnees est nommee "
    "investbuddy et est configuree avec un utilisateur dedie. Un script d'initialisation nomme init.sql est execute au premier "
    "demarrage pour creer les tables necessaires."
)

# Section 3: Volumes et persistance
doc.add_heading('3. Volumes et persistance des donnees', level=1)

doc.add_paragraph(
    "L'architecture utilise deux types de volumes pour assurer la persistance des donnees. Les volumes nommes mysql_data et n8n_data "
    "permettent de conserver respectivement les donnees de la base MySQL et les configurations de n8n entre les redemarrages des conteneurs."
)

doc.add_paragraph(
    "Des bind mounts sont egalement utilises pour partager des fichiers entre l'hote et les conteneurs. Le dossier ./data est monte "
    "dans le conteneur de prediction API pour permettre l'acces aux fichiers de donnees de reference et de production. Le dossier "
    "artifacts est partage pour que l'API puisse charger les modeles ML. Le dossier de l'agent est monte pour permettre le hot-reload "
    "pendant le developpement. Enfin, les workflows n8n sont partages pour que l'orchestrateur puisse charger le workflow predefini."
)

doc.add_paragraph(
    "Le fichier ref_data.csv contient les donnees de reference utilisees pour l'entrainement initial du modele. Le fichier prod_data.csv "
    "stocke les retours utilisateurs collectes via le mecanisme de feedback. Ces deux fichiers sont essentiels au fonctionnement "
    "du systeme de monitoring et de reentrainement automatique."
)

# Section 4: Choix technologiques
doc.add_heading('4. Choix technologiques', level=1)

doc.add_paragraph(
    "FastAPI a ete choisi pour developper les APIs en raison de sa performance, de sa documentation automatique et de son support "
    "natif des operations asynchrones. Ce framework est particulierement adapte aux applications ML qui necessitent des temps "
    "de reponse rapides pour le serving des modeles."
)

doc.add_paragraph(
    "ChromaDB a ete retenu comme base de donnees vectorielle pour l'agent RAG en raison de sa simplicite d'utilisation et de sa "
    "persistance native. Contrairement a des solutions comme Pinecone qui necessitent une connexion cloud, ChromaDB fonctionne "
    "entierement en local, ce qui correspond aux besoins du projet."
)

doc.add_paragraph(
    "OpenAI Embeddings avec le modele text-embedding-3-small a ete choisi pour la vectorisation des documents financiers. Ce modele "
    "offre un bon equilibre entre qualite des embeddings et cout computationnel. L'API OpenAI est egalement utilisee pour la generation "
    "des messages d'alerte personnalises dans le workflow n8n."
)

doc.add_paragraph(
    "n8n a ete prefere a d'autres orchestrateurs comme Zapier ou Make en raison de son caractere open source et de la possibilite "
    "de l'heberger localement. Cette solution permet un controle total sur les donnees et les workflows sans dependre d'un service tiers."
)

doc.add_paragraph(
    "RandomForest a ete choisi comme algorithme de prediction pour sa robustesse et sa capacite a generaliser sur des donnees "
    "financieres bruitees. Ce modele offre egalement l'avantage d'etre facilement serialisable et interpretable, ce qui est important "
    "dans un contexte de monitoring MLOps."
)

# Section 5: Pipeline ML
doc.add_heading('5. Pipeline Machine Learning', level=1)

doc.add_paragraph(
    "Le pipeline ML commence par la collecte de donnees historiques via l'API Binance. Les donnees recuperees comprennent les prix "
    "OHLCV (Open, High, Low, Close, Volume) pour les paires BTCUSDT et ETHUSDT sur un intervalle de 6 heures depuis janvier 2023."
)

doc.add_paragraph(
    "L'etape de feature engineering calcule plusieurs indicateurs techniques. Le RSI (Relative Strength Index) sur 14 periodes "
    "mesure la force relative du mouvement des prix. L'ATR (Average True Range) sur 14 periodes quantifie la volatilite. "
    "Les moyennes mobiles simples et exponentielles (SMA_20 et EMA_50) captent les tendances. Le taux de variation du volume "
    "complete ces indicateurs pour detecter les anomalies de marche."
)

doc.add_paragraph(
    "Les cibles predictives sont de deux types. La premiere est la volatilite future calculee comme l'ecart-type des rendements "
    "logarithmiques sur une fenetre de 4 periodes (soit 24 heures). La seconde est la direction du prix, une variable binaire "
    "indiquant si le prix va monter ou descendre."
)

doc.add_paragraph(
    "Les modeles sont entraînes sur les donnees normalisees avec un StandardScaler puis exportes au format pickle dans le dossier "
    "artifacts. Un fichier de reference ref_data.csv est genere pour permettre la detection de data drift en production."
)

# Section 6: Monitoring avec Evidently
doc.add_heading('6. Monitoring et detection de drift', level=1)

doc.add_paragraph(
    "Le monitoring du modele est assure par Evidently AI, un outil open source specialise dans l'analyse de drift et la qualite "
    "des donnees ML. Le dashboard Evidently compare en permanence les donnees de reference stockees dans ref_data.csv avec les "
    "donnees de production accumulees dans prod_data.csv."
)

doc.add_paragraph(
    "Les metriques suivies comprennent les indicateurs de data drift qui detectent les changements dans la distribution des features "
    "d'entree. Les metriques de classification comme le F1-score, la precision, le rappel et l'accuracy mesurent la performance "
    "du modele sur les donnees de production annotees par le feedback utilisateur."
)

doc.add_paragraph(
    "Le dossier reporting contient la configuration Evidently avec un Dockerfile dedie et un fichier project.py qui definit "
    "les rapports a generer. Le dashboard est accessible sur le port 8082 et permet de visualiser l'etat de sante du modele "
    "et d'identifier les necessites de reentrainement."
)

# Section 7: Human-in-the-Loop
doc.add_heading('7. Boucle Human-in-the-Loop', level=1)

doc.add_paragraph(
    "Le systeme Human-in-the-Loop permet de collecter le feedback des utilisateurs pour ameliorer continuellement le modele. "
    "Le processus se declenche lorsqu'un utilisateur clique sur le bouton de notification apres avoir recu une prediction. "
    "L'API de prediction appelle alors le webhook n8n avec les donnees de la prediction et les coordonnees de l'utilisateur."
)

doc.add_paragraph(
    "Le workflow n8n verifie d'abord si la volatilite predite depasse un seuil de 0.02, ce qui qualifie une volatilite elevee. "
    "Si c'est le cas, le workflow genere un message personnalise avec GPT-3.5 et envoie un email HTML formatte a l'utilisateur. "
    "Cet email presente la prediction et propose deux boutons : confirmer ou corriger."
)

doc.add_paragraph(
    "Lorsque l'utilisateur clique sur un des boutons, l'endpoint /feedback de l'API de prediction est appele avec les donnees "
    "de la prediction et le label de feedback. Ces informations sont normalisees avec le scaler existant puis stockees dans "
    "le fichier prod_data.csv. Chaque ligne de ce fichier contient les features normalisees, la prediction et le label de feedback "
    "qui represente la verite terrain selon l'utilisateur."
)

doc.add_paragraph(
    "Le feedback collecte sert a deux fins. Premierement, il alimente le monitoring Evidently pour detecter les ecarts entre "
    "les predictions et la realite. Deuxiemement, il constitue un jeu de donnees supplementaire pour le reentrainement du modele."
)

# Section 8: Reentrainement automatique
doc.add_heading('8. Reentrainement automatique', level=1)

doc.add_paragraph(
    "Le systeme implemente un mecanisme de reentrainement automatique base sur la quantite de feedback collecte. Un seuil K "
    "defini par la variable d'environnement RETRAIN_THRESHOLD (par defaut 10) determine quand declencher un reentrainement."
)

doc.add_paragraph(
    "A chaque reception de feedback, l'API verifie si le nombre de lignes dans prod_data.csv est un multiple de K. Si c'est le cas, "
    "le processus de reentrainement est lance en arriere-plan. Ce processus concatene les donnees de reference et les donnees "
    "de production pour creer un nouveau jeu d'entrainement enrichi."
)

doc.add_paragraph(
    "Un nouveau scaler est ajuste sur l'ensemble des donnees combinees, puis de nouveaux modeles RandomForest sont entraînes "
    "pour predire la volatilite et la direction. Les artefacts mis a jour remplacent les anciens fichiers pickle dans le dossier "
    "artifacts, et les variables globales de l'API sont mises a jour pour utiliser immediatement les nouveaux modeles."
)

doc.add_paragraph(
    "Ce mecanisme garantit que le modele s'adapte progressivement aux nouvelles conditions du marche sans intervention manuelle. "
    "L'approche par seuil de feedback permet de controler le compromis entre stabilite du modele et capacite d'adaptation."
)

# Section 9: Flux de donnees
doc.add_heading('9. Flux de donnees', level=1)

doc.add_paragraph(
    "Le flux de prediction commence par la saisie d'une requete utilisateur dans l'interface web. La webapp envoie les donnees "
    "de marche a l'API de prediction sur l'endpoint /predict. L'API normalise les donnees avec le scaler charge, applique le modele "
    "et retourne la prediction de volatilite et de direction. L'interface affiche alors le resultat a l'utilisateur."
)

doc.add_paragraph(
    "Le flux de notification se declenche quand l'utilisateur clique sur notifier. L'API de prediction prepare les URLs de feedback "
    "avec les parametres codes, puis appelle le webhook n8n. n8n genere le message et envoie l'email. L'utilisateur recoit l'email "
    "et clique sur un lien de feedback. L'API enregistre alors le retour dans prod_data.csv et declenche eventuellement un reentrainement."
)

doc.add_paragraph(
    "Le flux de consultation RAG permet a l'utilisateur de poser des questions financieres. La webapp envoie la question a l'API agent "
    "sur l'endpoint /ask. L'API agent recherche les passages pertinents dans ChromaDB puis utilise le LLM pour generer une reponse "
    "contextualisee. La reponse est retournee a l'interface pour affichage avec un effet typewriter."
)

# Save document
doc.save('Rapport_Architecture_InvestBuddy.docx')
print("Rapport genere avec succes : Rapport_Architecture_InvestBuddy.docx")